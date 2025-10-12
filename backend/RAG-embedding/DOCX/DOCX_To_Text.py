# docx_to_text.py

import os
from docx import Document
from PIL import Image
import pytesseract

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\tesseract.exe"  # Windows path

def docx_to_text(docx_path, output_dir="Text_files", output_name=None):
    """
    Extract text from DOCX files including:
    - Paragraphs
    - Tables
    - Text inside images/diagrams/charts (via OCR)
    
    Saves output as a .txt file in Text_files/ with the same original filename.
    """
    os.makedirs(output_dir, exist_ok=True)

    base_name = output_name if output_name else os.path.splitext(os.path.basename(docx_path))[0]
    output_txt_path = os.path.join(output_dir, f"{base_name}.txt")

    doc = Document(docx_path)
    all_text = ""

    # 1️⃣ Extract paragraphs
    for para_num, para in enumerate(doc.paragraphs, start=1):
        if para.text.strip():
            all_text += f"[Paragraph {para_num}]: {para.text.strip()}\n"

    # 2️⃣ Extract tables
    for table_num, table in enumerate(doc.tables, start=1):
        all_text += f"\n[Table {table_num}]:\n"
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            all_text += "\t".join(row_text) + "\n"

    # 3️⃣ Extract images (diagrams/charts)
    images_dir = os.path.join(output_dir, f"{base_name}_images")
    os.makedirs(images_dir, exist_ok=True)

    for i, rel in enumerate(doc.part._rels):
        rel = doc.part._rels[rel]
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image_path = os.path.join(images_dir, f"{base_name}_image_{i+1}.png")
            with open(image_path, "wb") as img_file:
                img_file.write(image_data)

            # OCR for text in images
            img = Image.open(image_path)
            text_in_image = pytesseract.image_to_string(img, lang="eng")
            if text_in_image.strip():
                all_text += f"\n[Text in Image {i+1}]:\n{text_in_image.strip()}\n"

    # Save all extracted text
    with open(output_txt_path, "w", encoding="utf-8") as f:
        f.write(all_text.strip())

    print(f"\n✅ DOCX text extraction complete! Saved to: {os.path.abspath(output_txt_path)}")
    return output_txt_path
